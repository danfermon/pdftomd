
# C:\pdftomd\pdf_detector.py

import fitz # PyMuPDF
from markitdown import MarkItDown
import os
from docx import Document # python-docx

def is_digital_pdf(pdf_path: str, min_text_chars: int = 30, threshold: float = 0.7) -> bool:
    """
    Verifica se um PDF é predominantemente digital (texto extraível) ou escaneado (imagem).

    Args:
        pdf_path: Caminho completo para o arquivo PDF.
        min_text_chars: Mínimo de caracteres para uma página ser considerada "com texto".
        threshold: Proporção mínima de páginas com texto para classificar o PDF como digital.

    Returns:
        True se for digital, False se for escaneado (imagem).
    """
    try:
        doc = fitz.open(pdf_path)
        total_pages = len(doc)
        if total_pages == 0:
            return False

        pages_with_text = 0
        for page in doc:
            # Extrai texto simples
            text = page.get_text()
            if len(text.strip()) >= min_text_chars:
                pages_with_text += 1

        ratio = pages_with_text / total_pages
        return ratio >= threshold

    except Exception as e:
        print(f"Erro na detecção do PDF {pdf_path}: {e}")
        return False

def preprocess_docx_with_pagination(docx_path: str) -> str:
    """
    Lê um arquivo DOCX, detecta quebras de página manuais e insere marcadores '## PÁGINA X'.
    Retorna o caminho de um arquivo temporário modificado.
    """
    try:
        doc = Document(docx_path)
        page_count = 1
        
        # Inserir marcador da página 1 no início
        # Verifica se o documento tem parágrafos antes de tentar inserir
        if doc.paragraphs:
            doc.paragraphs[0].insert_paragraph_before(f"## PÁGINA {page_count}")
        else:
            # Se o documento estiver vazio, adiciona o primeiro parágrafo com o marcador
            doc.add_paragraph(f"## PÁGINA {page_count}")


        for para in doc.paragraphs:
            # Check for explicit page breaks within the paragraph's runs
            for run in para.runs:
                # Check for w:br with type="page" (manual page break)
                # or w:lastRenderedPageBreak (often indicates a page break)
                if '<w:br w:type="page"/>' in run._element.xml or '<w:lastRenderedPageBreak/>' in run._element.xml:
                    page_count += 1
                    # Insert the new page marker before the current paragraph
                    para.insert_paragraph_before(f"\n\n## PÁGINA {page_count}")
                    # Break from inner loop as we've handled this paragraph's page break
                    break 
        
        temp_path = docx_path.replace(".docx", "_paginated_temp.docx")
        doc.save(temp_path)
        return temp_path
    except Exception as e:
        print(f"Aviso: Falha ao processar paginação do DOCX ({e}). Usando arquivo original.")
        return docx_path

def extract_structured_markitdown(file_path: str, output_md: str):
    """
    Extrai o texto de qualquer arquivo suportado (PDF digital, DOCX, XLSX, etc.) 
    usando MarkItDown para preservar a estrutura.
    Em caso de falha no PDF, tenta um fallback simples com PyMuPDF.
    """
    if not os.path.exists(file_path):
        print(f"Erro: Arquivo não encontrado em {file_path}")
        return

    processing_path = file_path
    is_temp_docx = False

    # --- Pré-processamento para DOCX (Paginação) ---
    if file_path.lower().endswith('.docx'):
        print("Pré-processando DOCX para inserir marcadores de página...")
        processing_path = preprocess_docx_with_pagination(file_path)
        if processing_path != file_path:
            is_temp_docx = True

    # --- Tentativa 1: MarkItDown ---
    try:
        md_converter = MarkItDown()
        print(f"Iniciando conversão estruturada de {os.path.basename(file_path)} para Markdown (MarkItDown)...")
        
        result = md_converter.convert(processing_path)
        markdown_content = result.text_content
        
        # REMOVIDO: A verificação de tabelas estava forçando fallback em documentos de texto puro.
        # Confiamos no MarkItDown para gerar o texto estruturado corretamente.
        
        with open(output_md, 'w', encoding='utf-8') as f:
            f.write(markdown_content)
        print(f"Extração estruturada (MarkItDown) concluída: {output_md}")
        
    except Exception as e:
        print(f"Erro ao usar MarkItDown para extração: {e}")
        
        # --- Tentativa 2: Fallback PyMuPDF (Apenas para PDFs) ---
        if file_path.lower().endswith('.pdf'):
            print("Tentando fallback com PyMuPDF (layout) para PDF...")
            try:
                doc = fitz.open(file_path)
                with open(output_md, 'w', encoding='utf-8') as f:
                    for i, page in enumerate(doc):
                        # Usamos 'text' com o parâmetro 'layout' para tentar preservar o espaçamento
                        text = page.get_text("text", sort=True) 
                        f.write(f"\n\n## Página {i+1}\n\n{text.strip()}\n")
                print(f"Extração digital (Fallback PyMuPDF com layout) concluída: {output_md}")
            except Exception as e_f:
                print(f"Falha total na extração digital de PDF: {e_f}")
        else:
            print("Nenhum fallback disponível para este tipo de arquivo.")
            
    finally:
        # Limpeza do arquivo temporário DOCX
        if is_temp_docx and os.path.exists(processing_path):
            try:
                os.remove(processing_path)
            except Exception as e:
                print(f"Aviso: Não foi possível remover o arquivo temporário DOCX {processing_path}: {e}")

